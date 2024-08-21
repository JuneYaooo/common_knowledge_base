# read_files.py

import os
import re
import zipfile
import requests
import docx
import html2text
import requests
from goose3 import Goose
from goose3.text import StopWordsChinese
from bs4 import BeautifulSoup
import trafilatura
import chardet
import random
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF
import logging

# 设置日志记录器
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
handler = logging.StreamHandler()
formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
handler.setFormatter(formatter)
logger.addHandler(handler)
def extract_text_html2text(url, ignore_links=False, ignore_images=False, ignore_videos=False):
    """
    从指定 URL 提取网页正文内容

    参数:
    url (str): 要提取的网页 URL
    ignore_links (bool): 是否忽略链接，默认 False 保留链接
    ignore_images (bool): 是否忽略图像，默认 False 保留图像
    ignore_videos (bool): 是否忽略视频，默认 False 保留视频

    返回:
    str: 提取的网页正文内容
    """
    try:
        # 获取网页 HTML 内容
        response = requests.get(url)
        response.raise_for_status()  # 检查请求是否成功
        html_content = response.text

        # 使用 html2text 库提取正文
        h = html2text.HTML2Text()
        h.ignore_links = ignore_links
        h.ignore_images = ignore_images
        h.ignore_videos = ignore_videos
        text_content = h.handle(html_content)

        return text_content
    except Exception as e:
        logger.info(f"html2text 提取失败: {e}")
        return ""

def extract_content_goose(url):
    try:
        g = Goose({'stopwords_class': StopWordsChinese})
        article = g.extract(url=url)

        # 获取文章的 HTML
        html = article.raw_html

        # 使用 BeautifulSoup 解析 HTML
        soup = BeautifulSoup(html, 'html.parser')

        # 替换图片
        for img in soup.find_all('img'):
            img_src = img.get('src')
            if img_src:
                img.replace_with(f'[Image: {img_src}]')

        # 替换视频
        for video in soup.find_all('video'):
            video_src = video.get('src')
            if video_src:
                video.replace_with(f'[Video: {video_src}]')

        # 替换链接
        for a in soup.find_all('a'):
            link_href = a.get('href')
            if link_href:
                a.replace_with(f'[Link: {link_href}]')

        # 替换代码块
        for code in soup.find_all('pre'):
            code_text = code.get_text()
            if code_text:
                code.replace_with(f'[Code Block: {code_text}]')

        # 获取处理后的正文文本
        text = soup.get_text()

        # 清理多余的空行和空白符
        lines = text.splitlines()
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        cleaned_text = "\n".join(cleaned_lines)

        return cleaned_text
    except Exception as e:
        logger.info(f"Goose 提取失败: {e}")
        return ""

def extract_content_trafilatura(url):
    try:
        downloaded = trafilatura.fetch_url(url)
        content = trafilatura.extract(downloaded, include_images=True, include_links=True, include_formatting=True)
        return content
    except Exception as e:
        logger.info(f"Trafilatura 提取失败: {e}")
        return ""


def extract_text(url, priority=None):
    """
    从指定 URL 提取网页正文内容，按优先级依次使用 html2text, Goose, Trafilatura。
    如果优先级列表为空，则随机排序这三种方法。
    """
    # 方法映射
    methods = {
        'Trafilatura': extract_content_trafilatura,
        'html2text': extract_text_html2text,
        'Goose': extract_content_goose
    }
    
    # 如果优先级列表为空，则随机排序
    if not priority:
        priority = list(methods.keys())
        random.shuffle(priority)
    
    # 按照优先级尝试每种方法
    for method_name in priority:
        method = methods[method_name]
        text = method(url)
        if text and len(text) > 50 and "当前环境异常，完成验证后即可继续访问" not in text:
            return text

    # 如果都提取失败，返回 None
    return None


def unzip_file(zip_file_path):
    # 定义解压目标目录
    extract_to_path = os.path.dirname(zip_file_path)
    extract_path = zip_file_path.rsplit('.', 1)[0]
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        # 获取 ZIP 文件内的所有文件名
        zip_info_list = zip_ref.infolist()
        for zip_info in zip_info_list:
            try:
                # 尝试使用 cp437 编码
                filename = zip_info.filename.encode('cp437').decode('utf-8')
            except UnicodeDecodeError as e1:
                logger.info(f"Failed to decode {zip_info.filename} using cp437: {e1}")
                try:
                    # 尝试使用 gbk 编码
                    filename = zip_info.filename.encode('cp437').decode('gbk')
                except UnicodeDecodeError as e2:
                    logger.info(f"Failed to decode {zip_info.filename} using gbk: {e2}")
                    try:
                        # 如果 cp437 和 gbk 解码失败，检测编码并尝试用检测到的编码
                        detected_encoding = chardet.detect(zip_info.filename.encode('utf-8'))
                        encoding = detected_encoding['encoding']
                        logger.info(f"Detected encoding for {zip_info.filename}: {encoding}")
                        filename = zip_info.filename.encode('utf-8').decode(encoding)
                    except (UnicodeDecodeError, TypeError) as e3:
                        logger.info(f"Error encoding filename: {zip_info.filename}")
                        logger.info(e3)
                        continue

            # 修改 zip_info 的 filename
            zip_info.filename = filename
            
            # 解压单个文件
            zip_ref.extract(zip_info, extract_to_path)

    return extract_path


def list_files_in_folder(folder_path):
    """
    读取指定文件夹中的所有文件，返回文件路径列表。
    """
    file_paths = []

    # 遍历文件夹中的所有文件
    for filename in os.listdir(folder_path):
        filepath = os.path.join(folder_path, filename)

        # 检查文件是否是普通文件
        if os.path.isfile(filepath):
            file_paths.append(filepath)

    return file_paths


def read_file(file_path):
    """
    根据文件路径读取文件后缀名、去掉后缀的文件名和文件内容，并返回一个包含这些信息的字典。
    如果无法成功读取文件，返回一个空列表。
    """
    result = {}

    # 提取文件后缀名和文件名
    file_name = os.path.basename(file_path)
    file_name_without_extension, file_extension = os.path.splitext(file_name)

    try:
        # 首先尝试使用 UTF-8 编码读取
        with open(file_path, 'r', encoding='utf-8') as file:
            file_content = file.read()
    except UnicodeDecodeError:
        # 如果 UTF-8 读取失败，尝试使用 GBK 编码
        try:
            with open(file_path, 'r', encoding='gbk') as file:
                file_content = file.read()
        except Exception:
            # 如果读取仍然失败，返回空列表
            return []
    except Exception:
        # 如果出现其他异常，也返回空列表
        return []

    # 将后缀名、去掉后缀的文件名和文件内容添加到结果字典中
    result['file_extension'] = file_extension
    result['file_name'] = file_name
    result['file_content'] = file_content

    return result


def extract_md_title(md_content):
    """
    从Markdown内容中提取标题。
    
    参数:
    md_content (str): Markdown文件的内容。
    
    返回:
    str: 提取到的标题，如果提取失败则返回''。
    """
    # 匹配Markdown标题的正则表达式
    title_pattern = re.compile(r'^\s*#* (.+?)\s*$', re.MULTILINE)

    # 在内容中查找匹配的标题
    match = title_pattern.search(md_content)

    # 如果找到匹配的标题，返回提取到的标题；否则返回空字符串
    if match:
        return match.group(1)
    else:
        return ''
    

def split_markdown(text):
    # 临时替换论文引用和作者信息
    citation_pattern = re.compile(r'\[\d+\]')
    citations = citation_pattern.findall(text)
    for i, citation in enumerate(citations):
        text = text.replace(citation, f'__CITATION_{i}__')

    author_pattern = re.compile(r'\[@[^\]]+\]')
    authors = author_pattern.findall(text)
    for i, author in enumerate(authors):
        text = text.replace(author, f'__AUTHOR_{i}__')

    # 临时替换表格内容
    table_pattern = re.compile(r'\|.*\|')
    tables = table_pattern.findall(text)
    for i, table in enumerate(tables):
        text = text.replace(table, f'__TABLE_{i}__')

    # 将换行符替换为空格
    text = text.replace('\n', ' ')

    # 使用正则表达式找到句子结束符号（句号、问号、感叹号）
    sentences = re.split(r'(?<=[.!?])\s+', text)

    # 还原论文引用和作者信息
    for i, citation in enumerate(citations):
        sentences = [sentence.replace(f'__CITATION_{i}__', citation) for sentence in sentences]

    for i, author in enumerate(authors):
        sentences = [sentence.replace(f'__AUTHOR_{i}__', author) for sentence in sentences]

    # 还原表格内容
    for i, table in enumerate(tables):
        sentences = [sentence.replace(f'__TABLE_{i}__', table) for sentence in sentences]

    return sentences


def split_markdown_by_headings_old(markdown_text):
    # 使用正则表达式按照标题拆分Markdown文本
    headings = re.split(r'\n\s*#', markdown_text.strip())

    # 在每个标题前添加#（除了第一个标题）
    paragraphs = ['#' + heading if i > 0 else heading for i, heading in enumerate(headings)]

    return paragraphs

def split_docs_by_headings_old(text):
    # 定义标题的正则表达式模式（这里假设标题格式为“数字. 标题”）
    title_pattern = re.compile(r'^[一二三四五六七八九十]+、')

    sections = []
    current_section = []

    for line in text.split('\n'):
        # 检查当前行是否为新的大标题
        if title_pattern.match(line.strip()):
            # 如果当前部分不为空，保存并开始一个新部分
            if current_section:
                sections.append('\n'.join(current_section))
                current_section = []
        current_section.append(line.strip())

    # 添加最后一部分
    if current_section:
        sections.append('\n'.join(current_section))

    return sections

import re

def split_markdown_by_headings(markdown_text, max_token_length=1000):
    def split_paragraphs(text, max_length):
        """拆分段落，如果段落长度超过最大限制，按完整句子拆分"""
        sentences = re.split(r'(?<=[。！？])', text)
        paragraphs = []
        current_paragraph = ""

        for sentence in sentences:
            if len(current_paragraph) + len(sentence) > max_length:
                if current_paragraph.strip():
                    paragraphs.append(current_paragraph.strip())
                current_paragraph = sentence
            else:
                current_paragraph += sentence

        if current_paragraph.strip():
            paragraphs.append(current_paragraph.strip())

        return paragraphs

    def is_table_line(line):
        """判断行是否属于表格"""
        return line.startswith('|')

    def remove_image_paths(text):
        """移除Markdown中的图片路径"""
        return re.sub(r'!\[.*?\]\(.*?\)', '', text)

    def check_paragraph_length(paragraph, max_length):
        """检查段落长度是否超过最大限制的5倍"""
        if len(paragraph) > max_length * 5:
            print(f"段落实际字数：{len(paragraph)}，超过最大限制的5倍，已删除该段落。")
            return True
        return False

    # 使用正则表达式去除参考文献及其后的内容
    markdown_text = re.sub(r'参考文献.*?(\n{2,}|\n#|$)', '', markdown_text, flags=re.DOTALL)
    markdown_text = re.sub(r'(###? ?参 考 献 文.*?|\n参 考 献 文.*?)(\n{2,}|\n#|$)', '', markdown_text, flags=re.DOTALL)
    
    # 使用正则表达式按照标题拆分Markdown文本
    headings = re.split(r'\n\s*(?=#)', markdown_text.strip())

    # 处理每个段落，提取标题级别和内容
    result = []
    current_title_level1, current_title_level2, current_title_level3 = None, None, None
    current_content = []
    table_content = []

    def add_paragraphs_to_result(title_level1, title_level2, title_level3, content):
        """合并同一标题下的内容，并按最大长度拆分后添加到结果中"""
        text = '\n'.join(content).strip()
        if text:
            paragraphs = split_paragraphs(text, max_token_length)
            for para in paragraphs:
                if not check_paragraph_length(para, max_token_length):
                    result.append({
                        'title_level1': title_level1,
                        'title_level2': title_level2,
                        'title_level3': title_level3,
                        'content': para
                    })

    for paragraph in headings:
        paragraph = paragraph.strip()
        if paragraph == '':
            continue

        # 使用正则表达式匹配标题级别和内容
        match = re.match(r'(#*)(.*)', paragraph)
        title_level = len(match.group(1))
        title_content = match.group(2).strip()

        if title_level == 1:
            if current_content or table_content:
                add_paragraphs_to_result(current_title_level1, current_title_level2, current_title_level3, current_content)
                current_content = []
                if table_content:
                    result.append({
                        'title_level1': current_title_level1,
                        'title_level2': current_title_level2,
                        'title_level3': current_title_level3,
                        'content': '\n'.join(table_content)
                    })
                    table_content = []
            current_title_level1 = title_content
            current_title_level2, current_title_level3 = None, None
        elif title_level == 2:
            if current_content or table_content:
                add_paragraphs_to_result(current_title_level1, current_title_level2, current_title_level3, current_content)
                current_content = []
                if table_content:
                    result.append({
                        'title_level1': current_title_level1,
                        'title_level2': current_title_level2,
                        'title_level3': current_title_level3,
                        'content': '\n'.join(table_content)
                    })
                    table_content = []
            current_title_level2 = title_content
            current_title_level3 = None
        elif title_level == 3:
            if current_content or table_content:
                add_paragraphs_to_result(current_title_level1, current_title_level2, current_title_level3, current_content)
                current_content = []
                if table_content:
                    result.append({
                        'title_level1': current_title_level1,
                        'title_level2': current_title_level2,
                        'title_level3': current_title_level3,
                        'content': '\n'.join(table_content)
                    })
                    table_content = []
            current_title_level3 = title_content

        content_lines = paragraph[len(match.group(0)):].strip().split('\n')
        for line in content_lines:
            line = line.strip()
            if is_table_line(line):
                if current_content:
                    add_paragraphs_to_result(current_title_level1, current_title_level2, current_title_level3, current_content)
                    current_content = []
                table_content.append(line)
            else:
                if table_content:
                    result.append({
                        'title_level1': current_title_level1,
                        'title_level2': current_title_level2,
                        'title_level3': current_title_level3,
                        'content': '\n'.join(table_content)
                    })
                    table_content = []
                line = remove_image_paths(line)
                if line:
                    current_content.append(line)

    if current_content:
        add_paragraphs_to_result(current_title_level1, current_title_level2, current_title_level3, current_content)
    if table_content:
        result.append({
            'title_level1': current_title_level1,
            'title_level2': current_title_level2,
            'title_level3': current_title_level3,
            'content': '\n'.join(table_content)
        })

    return result

def split_text_by_length(text, max_token_length=1000):
    def split_sentences(text, max_length):
        """按句子拆分段落，如果段落长度超过最大限制，按完整句子拆分"""
        sentences = re.split(r'(?<=[。！？])', text)
        paragraphs = []
        current_paragraph = ""

        for sentence in sentences:
            if len(current_paragraph) + len(sentence) > max_length:
                if current_paragraph.strip():
                    paragraphs.append(current_paragraph.strip())
                current_paragraph = sentence
            else:
                current_paragraph += sentence

        if current_paragraph.strip():
            paragraphs.append(current_paragraph.strip())

        return paragraphs

    def check_paragraph_length(paragraph, max_length):
        """检查段落长度是否超过最大限制的5倍"""
        if len(paragraph) > max_length * 5:
            print(f"段落实际字数：{len(paragraph)}，超过最大限制的5倍，已删除该段落。")
            return True
        return False

    # 按单行或双行分割
    paragraphs = re.split(r'(\n{1,2})', text)  # 匹配单行或双行换行符
    combined_paragraphs = []
    current_chunk = ""

    for i in range(0, len(paragraphs), 2):  # 每两行进行处理，避免分隔符被误加入段落
        paragraph = paragraphs[i].strip()
        if not paragraph:
            continue

        if len(current_chunk) + len(paragraph) > max_token_length:
            # 如果当前组合内容超出最大长度，先拆分成更小的部分
            small_paragraphs = split_sentences(paragraph, max_token_length)
            for small_paragraph in small_paragraphs:
                if len(current_chunk) + len(small_paragraph) > max_token_length:
                    if current_chunk.strip() and not check_paragraph_length(current_chunk, max_token_length):
                        combined_paragraphs.append(current_chunk.strip())
                    current_chunk = small_paragraph
                else:
                    current_chunk += "\n" + small_paragraph
        else:
            current_chunk += "\n" + paragraph

    # 添加最后剩余的内容
    if current_chunk.strip() and not check_paragraph_length(current_chunk, max_token_length):
        combined_paragraphs.append(current_chunk.strip())

    return combined_paragraphs

def split_docs_by_headings(text, max_token_length=1000):
    def split_paragraphs(text, max_length):
        """拆分段落，如果段落长度超过最大限制，按完整句子拆分"""
        sentences = re.split(r'(?<=[。！？])', text)
        paragraphs = []
        current_paragraph = ""
        
        for sentence in sentences:
            if len(current_paragraph) + len(sentence) > max_length:
                if current_paragraph.strip():
                    paragraphs.append(current_paragraph.strip())
                current_paragraph = sentence
            else:
                current_paragraph += sentence
        
        if current_paragraph.strip():
            paragraphs.append(current_paragraph.strip())
        
        return paragraphs
    
    def check_paragraph_length(paragraph, max_length):
        """检查段落长度是否超过最大限制的5倍"""
        if len(paragraph) > max_length * 5:
            print(f"段落实际字数：{len(paragraph)}，超过最大限制的5倍，已删除该段落。")
            return True
        return False

    # 使用正则表达式去除参考文献及其后的内容
    text = re.sub(r'参考文献.*?(\n{2,}|\n#|$)', '', text, flags=re.DOTALL)
    text = re.sub(r'(###? ?参 考 献 文.*?|\n?参 考 献 文.*?|###? ?考 参 献 文.*?|\n?考 参 献 文.*?|###? ?考 参 文 献.*?|\n?考 参 文 献.*?|###? ?参 考 文 献.*?|\n?参 考 文 献.*?)(\n{2,}|\n#|$)', '', text, flags=re.DOTALL)
    
    # 定义标题的正则表达式模式
    title_pattern_1 = re.compile(r'^([一二三四五六七八九十]+)、(.*)')
    title_pattern_2 = re.compile(r'^(\d+)\.(.*)')

    result = []
    current_title_level1 = None
    current_title_level2 = None
    current_content = []
    table_content = []

    def add_paragraphs_to_result(title_level1, title_level2, content):
        """合并同一标题下的内容，并按最大长度拆分后添加到结果中"""
        text = '\n'.join(content).strip()
        if text:
            paragraphs = split_paragraphs(text, max_token_length)
            for para in paragraphs:
                if not check_paragraph_length(para, max_token_length):
                    result.append({
                        'title_level1': title_level1,
                        'title_level2': title_level2,
                        'content': para
                    })

    for section in text.split('\n'):
        section = section.strip()
        if section == '':
            continue
        
        # 检查当前行是否为新的大标题
        match_1 = title_pattern_1.match(section)
        match_2 = title_pattern_2.match(section)
        
        if match_1:
            if current_content or table_content:
                add_paragraphs_to_result(current_title_level1, current_title_level2, current_content)
                current_content = []
                if table_content:
                    result.append({
                        'title_level1': current_title_level1,
                        'title_level2': current_title_level2,
                        'content': '\n'.join(table_content)
                    })
                    table_content = []
            current_title_level1 = match_1.group(2).strip()
            current_title_level2 = None
        elif match_2:
            if current_content or table_content:
                add_paragraphs_to_result(current_title_level1, current_title_level2, current_content)
                current_content = []
                if table_content:
                    result.append({
                        'title_level1': current_title_level1,
                        'title_level2': current_title_level2,
                        'content': '\n'.join(table_content)
                    })
                    table_content = []
            current_title_level2 = match_2.group(2).strip()
        else:
            if section.startswith('|'):
                if current_content:
                    add_paragraphs_to_result(current_title_level1, current_title_level2, current_content)
                    current_content = []
                table_content.append(section)
            else:
                if table_content:
                    result.append({
                        'title_level1': current_title_level1,
                        'title_level2': current_title_level2,
                        'content': '\n'.join(table_content)
                    })
                    table_content = []
                if section:
                    current_content.append(section)
    
    if current_content:
        add_paragraphs_to_result(current_title_level1, current_title_level2, current_content)
    if table_content:
        result.append({
            'title_level1': current_title_level1,
            'title_level2': current_title_level2,
            'content': '\n'.join(table_content)
        })

    return result
def split_paragraph_into_sentences(paragraph):
    # 使用中文句子结束符进行分割
    sentences = re.split(r'(?<=[。？！])', paragraph)
    # 移除空字符串
    sentences = [sentence.strip() for sentence in sentences if sentence.strip()]
    return sentences

def process_markdown_old(markdown_text, file_name):
    # 调用函数将Markdown拆分成段落
    paragraphs = split_markdown_by_headings(markdown_text)
    title = extract_md_title(markdown_text)

    # 构建结果列表
    result_list = []
    for idx, paragraph in enumerate(paragraphs, 1):
        # 将段落拆分成句子
        sentences = split_markdown(paragraph)
        for sentence in sentences:
            # 构建字典
            result_dict = {
                'text': sentence,
                '段落': paragraph,
                '文件名': file_name,
                '论文名':title,
            }

            # 添加到结果列表
            result_list.append(result_dict)
            # 添加文件名提问
            result_dict = {
                'text': '文献：'+file_name+' '+sentence,
                '段落': paragraph,
                '文件名': file_name,
                '论文名':title,
            }
            result_list.append(result_dict)
        # 添加文件名，所有段落
        result_dict = {
            'text': file_name,
            '段落': paragraph,
            '文件名': file_name,
            '论文名':title,
        }
        result_list.append(result_dict)

    return result_list


def is_text_garbled(text):
    # 中文与英文的简单乱码判断
    chinese_characters = re.findall(r'[\u4e00-\u9fff]', text)
    symbol_characters = re.findall(r'[\u0000-\u0020\u3000\uFFFD]', text)  # 各类符号和空白字符

    if len(chinese_characters) > 0:
        chinese_ratio = len(chinese_characters) / max(len(text), 1)
        symbol_ratio = len(symbol_characters) / max(len(text), 1)
        return chinese_ratio < 0.2 or symbol_ratio > 0.3

    non_ascii_ratio = sum(1 for char in text if ord(char) > 127) / max(len(text), 1)
    return non_ascii_ratio > 0.3

def extract_pdf(filepath):
    filename = os.path.basename(filepath)
    result = {'file_extension': 'pdf', 'file_name': filename}

    def clean_text(text):
        # 尝试先用 UTF-8 解码，如果失败则使用 GBK 解码
        try:
            return text.encode('utf-8', 'ignore').decode('utf-8')
        except UnicodeDecodeError:
            try:
                return text.encode('utf-8', 'ignore').decode('gbk')
            except UnicodeDecodeError:
                return text  # 如果两种解码都失败，保持原始文本

    # 0. 尝试使用 OCR API
    upload_url = os.getenv("OCR_API_URL")
    if upload_url:
        try:
            files = {'file': (filename, open(filepath, 'rb'))}
            response = requests.post(upload_url, data={'return_format': 'md'}, files=files)
            assert response.status_code == 200
            response_json = response.json()
            ocr_text = clean_text("".join(response_json.get('result', [])))
            logger.info(f"OCR API extracted: {ocr_text[:50]}...")  # 打印前 50 个字符
            if ocr_text and not is_text_garbled(ocr_text):
                result['file_content'] = ocr_text
                return result
            logger.info(f"OCR result seems garbled for {filename}")
        except Exception as e:
            logger.info(f"OCR API failed for {filename}: {e}")
            logger.info("status_code:", response.status_code if 'response' in locals() else "N/A")
            logger.info(response.text if 'response' in locals() else "N/A")

    # 1. 尝试使用 PyMuPDF (fitz)
    try:
        document = fitz.open(filepath)
        content = []
        for page in document:
            content.append(clean_text(page.get_text()))
        combined_text = clean_text("".join(content))
        logger.info(f"PyMuPDF extracted: {combined_text[:50]}...")  # 打印前 50 个字符
        if combined_text and not is_text_garbled(combined_text):
            result['file_content'] = combined_text
            return result
    except Exception as e:
        logger.info(f"PyMuPDF (fitz) failed for {filename}: {e}")

    # 2. 尝试使用 pdfplumber
    try:
        with pdfplumber.open(filepath) as pdf:
            content = []
            for page in pdf.pages:
                content.append(clean_text(page.extract_text()))
        combined_text = clean_text("".join(content))
        logger.info(f"pdfplumber extracted: {combined_text[:50]}...")  # 打印前 50 个字符
        if combined_text and not is_text_garbled(combined_text):
            result['file_content'] = combined_text
            return result
    except Exception as e:
        logger.info(f"pdfplumber failed for {filename}: {e}")

    # 3. 尝试使用 PyPDF2
    try:
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            content = []
            for page in reader.pages:
                content.append(clean_text(page.extract_text()))
        combined_text = clean_text("".join(content))
        logger.info(f"PyPDF2 extracted: {combined_text[:50]}...")  # 打印前 50 个字符
        if combined_text and not is_text_garbled(combined_text):
            result['file_content'] = combined_text
            return result
    except Exception as e:
        logger.info(f"PyPDF2 failed for {filename}: {e}")

    # 如果所有方法都失败或提取的内容为乱码，返回空
    logger.info(f"All extraction methods failed or text is garbled for {filename}")
    return {}


def read_docx(filepath):
    """
    从.docx文件中读取文本内容。

    参数:
    file_path (str): .docx文件的路径。

    返回:
    str: 文件中的文本内容。
    """
    try:
        filename = os.path.basename(filepath)
        result = {'file_extension':'doc', 'file_name':filename}
        doc = docx.Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        full_text = [item for item in full_text if item != '']
        result['file_content'] = '\n'.join(full_text)
        return result
    except Exception as e:
        logger.info(f"Error reading {filepath}: {e}")
        return ''

def process_docx_old(markdown_text, file_name):
    # 调用函数将Markdown拆分成段落
    paragraphs = split_docs_by_headings(markdown_text)
    title = paragraphs[0].split('\n')[0]

    # 构建结果列表
    result_list = []
    for idx, paragraph in enumerate(paragraphs, 1):
        # 将段落拆分成句子
        sentences = split_paragraph_into_sentences(paragraph)
        for sentence in sentences:
            # 构建字典
            result_dict = {
                'text': sentence,
                '段落': paragraph,
                '文件名': file_name,
                '论文名':title,
            }

            # 添加到结果列表
            result_list.append(result_dict)
            # 添加文件名提问
            result_dict = {
                'text': '文献：'+file_name+' '+sentence,
                '段落': paragraph,
                '文件名': file_name,
                '论文名':title,
            }
            result_list.append(result_dict)
        # 添加文件名，所有段落
        result_dict = {
            'text': file_name,
            '段落': paragraph,
            '文件名': file_name,
            '论文名':title,
        }
        result_list.append(result_dict)

    return result_list


def process_markdown(markdown_text, file_name):
    # 调用函数将Markdown拆分成段落
    paragraphs = split_markdown_by_headings(markdown_text)
    title = extract_md_title(markdown_text)

    # 构建结果列表
    result_list = []
    for idx, paragraph in enumerate(paragraphs, 1):
        # 将段落拆分成句子
        sentences = split_markdown(paragraph['content'])
        for sentence in sentences:
            # 构建字典
            result_dict = {
                'text': sentence,
                '段落': paragraph['content'],
                '文件名': file_name,
                '论文名':title,
                'title_level1': paragraph['title_level1'] if 'title_level1' in paragraph else None,
                'title_level2': paragraph['title_level2'] if 'title_level2' in paragraph else None,
                'title_level3': paragraph['title_level3'] if 'title_level3' in paragraph else None,
                'serial_number':idx
            }

            # 添加到结果列表
            result_list.append(result_dict)
            # 添加文件名提问
            result_dict = {
                'text': '文献：'+file_name+' '+sentence,
                '段落': paragraph['content'],
                '文件名': file_name,
                '论文名':title,
                'title_level1': paragraph['title_level1'] if 'title_level1' in paragraph else None,
                'title_level2': paragraph['title_level2'] if 'title_level2' in paragraph else None,
                'title_level3': paragraph['title_level3'] if 'title_level3' in paragraph else None,
                'serial_number':idx
            }
            result_list.append(result_dict)
        # 添加标题
        for subtitle in ['title_level1','title_level2','title_level3']:
            if subtitle in paragraph and paragraph[subtitle] is not None:
                result_dict = {
                    'text': paragraph[subtitle],
                    '段落': paragraph['content'],
                    '文件名': file_name,
                    '论文名':title,
                    'title_level1': paragraph['title_level1'] if 'title_level1' in paragraph else None,
                    'title_level2': paragraph['title_level2'] if 'title_level2' in paragraph else None,
                    'title_level3': paragraph['title_level3'] if 'title_level3' in paragraph else None,
                    'serial_number':idx
                }
        result_list.append(result_dict)
    # 添加文件名，所有段落
    result_dict = {
        'text': file_name,
        '段落': markdown_text,
        '文件名': file_name,
        '论文名':title
    }
    result_list.append(result_dict)

    return result_list

def process_docx(markdown_text, file_name):
    # 调用函数将Markdown拆分成段落
    paragraphs = split_docs_by_headings(markdown_text)
    logger.info('file_name',file_name)
    base_name, _ = os.path.splitext(file_name)
    logger.info('base_name',base_name)
    title = paragraphs[0]['content'] if len(paragraphs)>0 and 'content' in paragraphs[0] else base_name

    # 构建结果列表
    result_list = []
    for idx, paragraph in enumerate(paragraphs, 1):
        # 将段落拆分成句子
        sentences = split_paragraph_into_sentences(paragraph['content'])
        for sentence in sentences:
            # 构建字典
            result_dict = {
                'text': sentence,
                '段落': paragraph['content'],
                '文件名': file_name,
                '论文名':title,
                'title_level1': paragraph['title_level1'] if 'title_level1' in paragraph else None,
                'title_level2': paragraph['title_level2'] if 'title_level2' in paragraph else None,
                'title_level3': paragraph['title_level3'] if 'title_level3' in paragraph else None,
                'serial_number':idx
            }
            

            # 添加到结果列表
            result_list.append(result_dict)
            # 添加文件名提问
            result_dict = {
                'text': '文献：'+file_name+' '+sentence,
                '段落': paragraph['content'],
                '文件名': file_name,
                '论文名':title,
                'title_level1': paragraph['title_level1'] if 'title_level1' in paragraph else None,
                'title_level2': paragraph['title_level2'] if 'title_level2' in paragraph else None,
                'title_level3': paragraph['title_level3'] if 'title_level3' in paragraph else None,
                'serial_number':idx
            }
            result_list.append(result_dict)
        # 添加标题
        for subtitle in ['title_level1','title_level2','title_level3']:
            if subtitle in paragraph and paragraph[subtitle] is not None:
                result_dict = {
                    'text': paragraph[subtitle],
                    '段落': paragraph['content'],
                    '文件名': file_name,
                    '论文名':title,
                    'title_level1': paragraph['title_level1'] if 'title_level1' in paragraph else None,
                    'title_level2': paragraph['title_level2'] if 'title_level2' in paragraph else None,
                    'title_level3': paragraph['title_level3'] if 'title_level3' in paragraph else None,
                    'serial_number':idx
                }
                result_list.append(result_dict)
        result_dict = {
        'text': file_name,
        '段落': markdown_text,
        '文件名': file_name,
        '论文名':title
    }
        result_list.append(result_dict)
    # 添加文件名，所有段落
    result_dict = {
        'text': file_name,
        '段落': markdown_text,
        '文件名': file_name,
        '论文名':title
    }
    result_list.append(result_dict)

    return result_list

def process_txt(txt_text, file_name):
    # 调用函数将txt文本拆分成段落
    paragraphs = split_text_by_length(txt_text)
    title = file_name  # 假设没有单独的标题，使用文件名作为标题

    # 构建结果列表
    result_list = []
    for idx, paragraph in enumerate(paragraphs, 1):
        # 将段落拆分成句子
        sentences = re.split(r'(?<=[。！？])', paragraph)
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            # 构建字典
            result_dict = {
                'text': sentence,
                '段落': paragraph,
                '文件名': file_name,
                '论文名': title,
                'serial_number': idx
            }

            # 添加到结果列表
            result_list.append(result_dict)

            # 添加带文件名提问的版本
            result_dict = {
                'text': '文献：' + file_name + ' ' + sentence,
                '段落': paragraph,
                '文件名': file_name,
                '论文名': title,
                'serial_number': idx
            }
            result_list.append(result_dict)

    # 添加整篇文章信息
    result_dict = {
        'text': file_name,
        '段落': txt_text,
        '文件名': file_name,
        '论文名': title
    }
    result_list.append(result_dict)

    return result_list
def process_url(markdown_text, url):
    # 调用函数将Markdown拆分成段落
    paragraphs = split_markdown_by_headings(markdown_text)
    logger.info('文件名',url)

    # 构建结果列表
    result_list = []
    for idx, paragraph in enumerate(paragraphs, 1):
        # 将段落拆分成句子
        sentences = split_paragraph_into_sentences(paragraph['content'])
        for sentence in sentences:
            # 构建字典
            result_dict = {
                'text': sentence,
                '段落': paragraph['content'],
                '文件名': url,
                'title_level1': paragraph['title_level1'] if 'title_level1' in paragraph else None,
                'title_level2': paragraph['title_level2'] if 'title_level2' in paragraph else None,
                'title_level3': paragraph['title_level3'] if 'title_level3' in paragraph else None,
                'serial_number':idx
            }
            

            # 添加到结果列表
            result_list.append(result_dict)
    # 添加url，所有段落
    result_dict = {
        'text': url,
        '段落': markdown_text,
    }
    result_list.append(result_dict)

    return result_list
def process_and_deduplicate(data_list):
    # 去掉每个字典中的'text'字段
    processed_list = [{key: value for key, value in d.items() if key != 'text'} for d in data_list]

    # 对处理后的列表进行字典去重，确保'serial_number'不重复
    unique_list = [dict(t) for t in {tuple(d.items()) for d in processed_list}]

    # 检查是否包含'serial_number'键，然后按'serial_number'从小到大排序
    sorted_list = sorted(unique_list, key=lambda x: x.get('serial_number', 0))
    return sorted_list


def read_url(url):
    """
    从.docx文件中读取文本内容。

    参数:
    file_path (str): .docx文件的路径。

    返回:
    str: 文件中的文本内容。
    """
    try:
        result = {'file_extension':'url', 'file_name':url}
        result['file_content'] = extract_text(url)
        if "当前环境异常，完成验证后即可继续访问" in result['file_content'] or ("参数错误" in result['file_content'] and len(result['file_content'])<60):
            return ''
        else:
            return result
    except Exception as e:
        logger.info(f"Error reading {url}: {e}")
        return ''
